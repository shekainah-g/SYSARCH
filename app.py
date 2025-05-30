from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, send_from_directory
import dbhelper
from dbhelper import get_db_connection
import os
from flask import jsonify
import time
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from docx import Document
from openpyxl import Workbook
import io
from werkzeug.utils import secure_filename
from flask_login import login_required, current_user, LoginManager, UserMixin
from flask_sqlalchemy import SQLAlchemy

app = Flask(__name__)
app.secret_key = 'g@ceta' 

# SQLAlchemy configuration
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# Flask-Login configuration
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

UPLOAD_FOLDER = "static/uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Add this to your database models section
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(120), nullable=False)
    firstname = db.Column(db.String(80))
    lastname = db.Column(db.String(80))
    midname = db.Column(db.String(80))
    email = db.Column(db.String(120))
    role = db.Column(db.String(20))
    course = db.Column(db.String(80))
    yearlvl = db.Column(db.String(20))
    profile_pic = db.Column(db.String(200))
    remaining_sessions = db.Column(db.Integer, default=0)
    notifications = db.relationship('Notification', backref='user', lazy=True)

class Notification(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    message = db.Column(db.String(255), nullable=False)
    read = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# Create all database tables
with app.app_context():
    db.create_all()

@app.route('/')
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if not username or not password:
            flash('Please enter both username and password', 'error')
            return redirect(url_for('login'))
        
        user = dbhelper.get_users(username, password)
        
        if user:
            session['username'] = user['data']['username']
            session['role'] = user['role']
            if user['role'] == 'student':
                session['user_id'] = user['data']['id']
            
            if user['role'] == 'admin':
                return redirect(url_for('admin_dashboard'))
            else:
                return redirect(url_for('student_dashboard'))
        else:
            flash('Invalid username or password', 'error')
            return redirect(url_for('login'))
    
    return render_template('login.html')

@app.route('/login')
def show_login():
    return render_template("login.html")

@app.route("/logout", methods=["POST"])
def logout():
    """Secure logout route using POST request"""
    session.clear()  # Clears all session data
    return redirect(url_for("login"))  # Redirect to the login page


@app.route('/dashboard')
def dashboard():
    if 'username' not in session:
        flash('Please log in first.', 'error')
        return redirect(url_for('login'))

    user = dbhelper.get_user_info(session['username'])

    if not user:
        flash("User not found!", "error")
        return redirect(url_for("login"))

    return render_template('dashboard.html', user=user)

@app.route('/admin_dashboard')
def admin_dashboard():
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))
    
    statistics = dbhelper.get_statistics()
    announcements = dbhelper.get_announcements()
    top_students = dbhelper.get_top_students()
    daily_stats = dbhelper.get_daily_sitins_stats()
    
    return render_template('admin_dashboard.html', 
                         statistics=statistics, 
                         announcements=announcements,
                         top_students=top_students,
                         daily_stats=daily_stats)

@app.route('/student_dashboard')
def student_dashboard():
    if 'username' not in session or session['role'] != 'student':
        return redirect(url_for('login'))
    
    return redirect(url_for('information'))

@app.route('/information')
def information():
    if 'username' not in session:
        flash('Please log in first.', 'error')
        return redirect(url_for('login'))  

    user = dbhelper.get_user_info(session['username'])  

    if not user:
        flash("User not found!", "error")
        return redirect(url_for("login"))

    # Get the student's points
    points = dbhelper.get_student_points(user['id'])
    user['points'] = points

    # Get announcements for display
    announcements = dbhelper.get_announcements()

    return render_template('information.html', user=user, announcements=announcements)



@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        id = request.form.get('id')
        password = request.form.get('password')
        firstname = request.form.get('firstname')
        lastname = request.form.get('lastname')
        midname = request.form.get('midname')
        email = request.form.get('email')
        username = request.form.get('username')
        course = request.form.get('course')
        yearlvl = request.form.get('yearlvl')
        
        if not all([id, username, password, firstname, lastname, midname, email, course, yearlvl]):
            flash('Please fill in all required fields', 'error')
            return redirect(url_for('register'))
        
        result = dbhelper.register_user(
            id, username, password, firstname, lastname, midname,
            email, 'student', course, yearlvl
        )
        
        if result == "Registration successful":
            flash('Registration successful! Please login.', 'success')
            return redirect(url_for('login'))
        else:
            flash(result, 'error')
            return redirect(url_for('register'))
    
    return render_template('register.html')

@app.route('/labrules')
def labrules():
    
    return render_template('labrules.html')


@app.route('/edit', methods=['GET', 'POST'])
def edit_profile():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    user = dbhelper.get_user_info(session['username'])  

    if request.method == 'POST':
        lastname = request.form['lastname']
        firstname = request.form['firstname']
        midname = request.form.get('midname', '')
        course = request.form['course']
        yearlvl = request.form['yearlvl']
        email = request.form['email']
        password = request.form.get('password', '')

        # Handling profile picture upload
        file = request.files.get("profile_pic")
        filename = user["profile_pic"]  # Default to the existing profile pic

        if file and file.filename:
            try:
                # Get the file extension
                file_ext = os.path.splitext(file.filename)[1]
                # Create a unique filename using username and timestamp
                filename = f"{session['username']}_{int(time.time())}{file_ext}"
                
                # Ensure the upload folder exists
                upload_folder = app.config["UPLOAD_FOLDER"]
                os.makedirs(upload_folder, exist_ok=True)
                
                # Save the file
                file_path = os.path.join(upload_folder, filename)
                file.save(file_path)
                print(f"File saved successfully to: {file_path}")  # Debug print
            except Exception as e:
                print(f"Error saving file: {str(e)}")  # Debug print
                flash('Error uploading profile picture. Please try again.', 'error')
                return redirect(url_for('edit_profile'))

        # Update the database with the new profile pic filename
        dbhelper.update_user(session['username'], lastname, firstname, midname, 
                           course, yearlvl, email, password, filename)

        flash('Profile updated successfully!', 'success')
        return redirect(url_for('information'))

    return render_template('edit.html', user=user)


# @app.route('/student_records', methods=['GET', 'POST'])
# def student_records():
#     if 'username' not in session:
#         flash('Please log in first.', 'error')
        
#     return render_template('student_records.html')
    
@app.route('/index')
def index():
    if 'username' not in session:
        flash('Please log in first.', 'warning')
        return redirect(url_for('login'))  

    return render_template('index.html') 

@app.route('/sitin')
def sitin():
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))
    
    current_sitins = dbhelper.get_current_sitins()
    return render_template('sitin.html', current_sitins=current_sitins)

@app.route('/admin')
def admin():
    return render_template('admin.html')


@app.route('/home')
def home():

    return render_template('home.html')


@app.route('/search', methods=['POST'])
def search():
    try:
        # Try to get query from form data first, then from JSON
        query = request.form.get('query', '').strip()
        if not query:
            query = request.json.get('query', '').strip()
            
        if not query:
            return jsonify({'error': 'Please enter a search query'})
        
        student = dbhelper.get_student_by_id_or_name(query)
        if student:
            return jsonify({
                'success': True,
                'student': {
                    'id': student['id'],
                    'name': student['name'],
                    'course': student['course'],
                    'yearlvl': student['yearlvl'],
                    'remaining_sessions': student['remaining_sessions']
                }
            })
        return jsonify({'error': 'Student not found'})
    except Exception as e:
        print(f"Error in search: {e}")
        return jsonify({'error': 'Error searching for student'})

@app.route('/log_in_student', methods=['POST'])
def log_in_student():
    try:
        student_id = request.form.get('student_id')
        purpose = request.form.get('purpose')
        laboratory = request.form.get('laboratory')
        
        if not all([student_id, purpose, laboratory]):
            return jsonify({'success': False, 'message': 'Missing required fields'})
        
        success, message = dbhelper.log_in_student(student_id, purpose, laboratory)
        return jsonify({'success': success, 'message': message})
    except Exception as e:
        print(f"Error in log_in_student route: {e}")
        return jsonify({'success': False, 'message': 'Error logging in student'})

@app.route('/log_out_student', methods=['POST'])
def log_out_student():
    try:
        data = request.get_json()
        student_id = data.get('student_id')
        
        if not student_id:
            return jsonify({'success': False, 'message': 'Missing student ID'})
        
        success, message = dbhelper.log_out_student(student_id)
        return jsonify({'success': success, 'message': message})
    except Exception as e:
        print(f"Error in log_out_student: {e}")
        return jsonify({'success': False, 'message': 'Error logging out student'})

@app.route('/get_current_sitins')
def get_current_sitins():
    try:
        sitins = dbhelper.get_current_sitins()
        return jsonify({'success': True, 'sitins': sitins})
    except Exception as e:
        print(f"Error in get_current_sitins: {e}")
        return jsonify({'error': 'Error fetching current sit-ins'})

@app.route('/get_sit_in_history')
def get_sit_in_history():
    try:
        history = dbhelper.get_sit_in_history()
        return jsonify({'success': True, 'history': history})
    except Exception as e:
        print(f"Error in get_sit_in_history: {e}")
        return jsonify({'error': 'Error fetching sit-in history'})

@app.route('/get_statistics')
def get_statistics():
    try:
        stats = dbhelper.get_statistics()
        return jsonify({'success': True, 'statistics': stats})
    except Exception as e:
        print(f"Error in get_statistics: {e}")
        return jsonify({'error': 'Error fetching statistics'})

@app.route('/post_announcement', methods=['POST'])
def post_announcement():
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized'})
    
    announcement = request.form.get('announcement')
    if not announcement:
        return jsonify({'success': False, 'message': 'Announcement cannot be empty'})
    
    date_created = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    success = dbhelper.create_announcement(announcement, date_created)
    
    if success:
        return jsonify({'success': True, 'message': 'Announcement posted successfully'})
    else:
        return jsonify({'success': False, 'message': 'Failed to post announcement'})

@app.route('/sit_in_history')
def sit_in_history():
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))
    
    history = dbhelper.get_sit_in_history()
    return render_template('sit_in_history.html', history=history)

@app.route('/search_student')
def search_student():
    try:
        query = request.args.get('query', '')
        if not query:
            return jsonify({"error": "No search query provided"}), 400
            
        student = dbhelper.get_student_by_id_or_name(query)
        if student:
            return jsonify(student)
        return jsonify({"error": "No student found"}), 404
    except Exception as e:
        print(f"Error in search_student: {str(e)}")
        return jsonify({"error": "Internal server error"}), 500

@app.route('/cs_history')
def cs_history():
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))
    
    history = dbhelper.get_today_sit_in_history()
    return render_template('cs_history.html', history=history)

@app.route('/get_today_history')
def get_today_history():
    try:
        history = dbhelper.get_today_sit_in_history()
        return jsonify({'success': True, 'history': history})
    except Exception as e:
        print(f"Error in get_today_history: {e}")
        return jsonify({'error': 'Error fetching today\'s history'})

@app.route('/student_list')
def student_list():
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))
    
    data = dbhelper.get_all_users()
    return render_template('student_list.html', 
                         users=data['users'],
                         year_stats=data['year_stats'],
                         course_stats=data['course_stats'],
                         purpose_stats=data['purpose_stats'])

@app.route('/reservation')
def reservation():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return render_template('reservation.html')

@app.route('/get_available_slots', methods=['POST'])
def get_available_slots():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    laboratory = data.get('laboratory')
    date = data.get('date')
    
    if not laboratory or not date:
        return jsonify({'error': 'Missing parameters'}), 400
    
    try:
        slots = dbhelper.get_available_time_slots(laboratory, date)
        if not slots:
            # If no slots are returned, generate default slots
            slots = [
                {"start": "08:00", "end": "09:00"},
                {"start": "09:00", "end": "10:00"},
                {"start": "10:00", "end": "11:00"},
                {"start": "11:00", "end": "12:00"},
                {"start": "13:00", "end": "14:00"},
                {"start": "14:00", "end": "15:00"},
                {"start": "15:00", "end": "16:00"},
                {"start": "16:00", "end": "17:00"}
            ]
        return jsonify({'slots': slots})
    except Exception as e:
        print(f"Error getting available slots: {e}")
        return jsonify({'error': 'Error getting available slots'}), 500

@app.route('/create_reservation', methods=['POST'])
def create_reservation():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        data = request.get_json()
        laboratory = data.get('laboratory')
        purpose = data.get('purpose')
        date = data.get('date')
        start_time = data.get('start_time')
        end_time = data.get('end_time')
        computer_id = data.get('computer_id')
        
        if not all([laboratory, purpose, date, start_time, end_time, computer_id]):
            return jsonify({'error': 'Missing parameters'}), 400
        
        success, message = dbhelper.create_reservation(
            session['user_id'],
            laboratory,
            purpose,
            date,
            start_time,
            end_time,
            computer_id
        )
        
        if success:
            return jsonify({'message': message})
        else:
            return jsonify({'error': message}), 400
    except Exception as e:
        print(f"Error in create_reservation: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/get_my_reservations')
@login_required
def get_my_reservations():
    try:
        if current_user.role != 'student':
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        reservations = dbhelper.get_student_reservations(current_user.id)
        
        # Convert to list of dicts
        reservations_list = []
        for r in reservations:
            reservations_list.append({
                'id': r[0],
                'laboratory': r[2],
                'computer_number': r[11] if len(r) > 11 else 'N/A',
                'reservation_date': r[4],
                'start_time': r[5],
                'end_time': r[6],
                'purpose': r[3],
                'status': r[8] if len(r) > 8 else 'pending'
            })
        
        return jsonify({
            'success': True,
            'reservations': reservations_list
        })
    except Exception as e:
        print(f"Error getting student reservations: {e}")
        return jsonify({
            'success': False,
            'error': str(e),
            'reservations': []
        })

@app.route('/admin/reservations')
def admin_reservations():
    try:
        # Get pending reservations with additional details
        reservations = dbhelper.get_pending_reservations()
        return render_template('admin_reservations.html', reservations=reservations)
    except Exception as e:
        print(f"Error loading admin reservations: {e}")
        return render_template('admin_reservations.html', reservations=[])

@app.route('/admin/get_reservations')
def get_reservations():
    try:
        laboratory = request.args.get('laboratory', '')
        reservations = dbhelper.get_pending_reservations()
        
        # Filter by laboratory if specified
        if laboratory:
            reservations = [r for r in reservations if r[2] == laboratory]  # Assuming index 2 is laboratory
        
        # Convert to list of dicts for JSON serialization
        reservations_list = []
        for r in reservations:
            reservations_list.append({
                'id': r[0],
                'user_id': r[1],
                'laboratory': r[2],
                'purpose': r[3],
                'reservation_date': r[4],
                'start_time': r[5],
                'end_time': r[6],
                'computer_number': r[11] if len(r) > 11 else 'N/A',
                'student_name': f"{r[9]} {r[10]}" if len(r) > 10 else 'Unknown',
                'course': r[12] if len(r) > 12 else '',
                'yearlvl': r[13] if len(r) > 13 else ''
            })
        
        return jsonify({
            'success': True,
            'reservations': reservations_list
        })
    except Exception as e:
        print(f"Error getting reservations: {e}")
        return jsonify({
            'success': False,
            'error': str(e),
            'reservations': []
        })

@app.route('/admin/update_reservation_status', methods=['POST'])
def update_reservation_status():
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        data = request.get_json()
        reservation_id = data.get('reservation_id')
        status = data.get('status')
        
        if not reservation_id or not status:
            return jsonify({'error': 'Missing required parameters'}), 400
        
        success, message = dbhelper.update_reservation_status(reservation_id, status)
        
        if success:
            return jsonify({'success': True, 'message': message})
        else:
            return jsonify({'error': message}), 400
    except Exception as e:
        print(f"Error updating reservation status: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/download_student_list/<format>')
def download_student_list(format):
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))
    
    data = dbhelper.get_all_users()
    users = data['users']
    
    if format == 'pdf':
        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        
        # Create table data
        table_data = [['Student ID', 'Name', 'Course', 'Year Level', 'Remaining Sessions']]
        for user in users:
            table_data.append([
                str(user['id']),
                f"{user['firstname']} {user['lastname']}",
                user['course'],
                user['yearlvl'],
                str(user['remaining_sessions'])
            ])
        
        # Create table
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='student_list.pdf', mimetype='application/pdf')
    
    elif format == 'docx':
        # Create Word document
        doc = Document()
        doc.add_heading('Student List', 0)
        
        # Create table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Add headers
        headers = ['Student ID', 'Name', 'Course', 'Year Level', 'Remaining Sessions']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        # Add data
        for user in users:
            row = table.add_row()
            row.cells[0].text = str(user['id'])
            row.cells[1].text = f"{user['firstname']} {user['lastname']}"
            row.cells[2].text = user['course']
            row.cells[3].text = user['yearlvl']
            row.cells[4].text = str(user['remaining_sessions'])
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='student_list.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    elif format == 'xlsx':
        # Create Excel workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Student List"
        
        # Add headers
        headers = ['Student ID', 'Name', 'Course', 'Year Level', 'Remaining Sessions']
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)
        
        # Add data
        for row, user in enumerate(users, 2):
            ws.cell(row=row, column=1, value=user['id'])
            ws.cell(row=row, column=2, value=f"{user['firstname']} {user['lastname']}")
            ws.cell(row=row, column=3, value=user['course'])
            ws.cell(row=row, column=4, value=user['yearlvl'])
            ws.cell(row=row, column=5, value=user['remaining_sessions'])
        
        # Save to buffer
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name='student_list.xlsx', mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    
    return redirect(url_for('student_list'))

@app.route('/student_history')
def student_history():
    if 'username' not in session:
        flash('Please log in first.', 'error')
        return redirect(url_for('login'))
    
    try:
        # Get user info first to get the user ID
        user = dbhelper.get_user_info(session['username'])
        if not user:
            flash('User not found', 'error')
            return redirect(url_for('login'))
            
        # Get student's history using the user ID
        history = dbhelper.get_student_history(user['id'])
        return render_template('student_history.html', history=history)
    except Exception as e:
        print(f"Error in student_history: {e}")
        flash('Error loading history', 'error')
        return redirect(url_for('information'))

@app.route('/submit_feedback', methods=['POST'])
def submit_feedback():
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Unauthorized'})
    
    try:
        data = request.get_json()
        sitin_id = data.get('sitin_id')
        feedback = data.get('feedback')
        
        if not sitin_id or not feedback:
            return jsonify({'success': False, 'message': 'Missing required fields'})
        
        success = dbhelper.add_feedback(sitin_id, session['user_id'], feedback)
        if success:
            return jsonify({'success': True, 'message': 'Feedback submitted successfully'})
        else:
            return jsonify({'success': False, 'message': 'Error submitting feedback'})
    except Exception as e:
        print(f"Error in submit_feedback: {e}")
        return jsonify({'success': False, 'message': 'Error submitting feedback'})

@app.route('/add_points', methods=['POST'])
def add_points():
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({"success": False, "message": "Unauthorized"})
    
    try:
        data = request.json
        student_id = data.get('student_id')
        points = data.get('points')
        
        if not student_id or not points:
            return jsonify({"success": False, "message": "Missing required data"})
        
        result = dbhelper.add_points(student_id, points)
        return jsonify(result)
        
    except Exception as e:
        print(f"Error in add_points: {e}")
        return jsonify({"success": False, "message": "Error adding points"})

@app.route('/get_available_computers')
def get_available_computers():
    if 'user_id' not in session:
        return jsonify({'error': 'Not logged in'}), 401
    
    laboratory = request.args.get('laboratory')
    date = request.args.get('date')
    start_time = request.args.get('start_time')
    end_time = request.args.get('end_time')
    
    if not all([laboratory, date, start_time, end_time]):
        return jsonify({'error': 'Missing required parameters'}), 400
    
    try:
        result = dbhelper.get_available_computers(laboratory, date, start_time, end_time)
        if result is None:
            return jsonify({'error': 'Error getting available computers'}), 500
        return jsonify(result)
    except Exception as e:
        print(f"Error in get_available_computers: {e}")
        return jsonify({'error': 'Error getting available computers'}), 500

@app.route('/lab_schedule')
def lab_schedule():
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))
    
    schedules = dbhelper.get_lab_schedules()
    return render_template('lab_schedule.html', schedules=schedules)

@app.route('/add_lab_schedule', methods=['POST'])
def add_lab_schedule():
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized'})
    
    try:
        laboratory = request.form['laboratory']
        days = request.form.getlist('days')  # Get list of selected days
        start_time = request.form['start_time']
        end_time = request.form['end_time']
        subject = request.form['subject']
        instructor = request.form['instructor']
        description = request.form.get('description')
        
        if not days:
            flash('Please select at least one day', 'error')
            return redirect(url_for('lab_schedule'))
        
        success, message = dbhelper.add_lab_schedule(
            laboratory, days, start_time, end_time,
            subject, instructor, description
        )
        
        if success:
            flash(message, 'success')
        else:
            flash(message, 'error')
            
        return redirect(url_for('lab_schedule'))
    except Exception as e:
        print(f"Error in add_lab_schedule: {e}")
        flash('Error adding schedule', 'error')
        return redirect(url_for('lab_schedule'))

@app.route('/delete_lab_schedule/<int:schedule_id>', methods=['DELETE'])
def delete_lab_schedule(schedule_id):
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized'})
    
    try:
        success = dbhelper.delete_lab_schedule(schedule_id)
        return jsonify({'success': success})
    except Exception as e:
        print(f"Error in delete_lab_schedule: {e}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/check_lab_availability')
def check_lab_availability():
    laboratory = request.args.get('laboratory')
    day_of_week = request.args.get('day_of_week')
    time = request.args.get('time')
    
    if not all([laboratory, day_of_week, time]):
        return jsonify({'available': False, 'message': 'Missing parameters'})
    
    is_available = dbhelper.is_lab_available(laboratory, day_of_week, time)
    return jsonify({'available': is_available})

@app.route('/reservation_history')
def reservation_history():
    if 'user_id' not in session:
        flash('Please log in first.', 'error')
        return redirect(url_for('login'))
    
    try:
        reservations = dbhelper.get_student_reservations(session['user_id'])
        return render_template('reservation_history.html', reservations=reservations)
    except Exception as e:
        print(f"Error in reservation_history: {e}")
        flash('Error loading reservation history', 'error')
        return redirect(url_for('dashboard'))

# Resources routes
@app.route('/resources')
def resources():
    if 'user_id' not in session:
        flash('Please log in first.', 'error')
        return redirect(url_for('login'))
    
    try:
        resources = dbhelper.get_resources()
        print("DEBUG: resources fetched for student:", resources)  # Debug print
        return render_template('resources.html', resources=resources)
    except Exception as e:
        print(f"Error: {e}")
        flash('Error loading resources', 'error')
        return redirect(url_for('resources'))

@app.route('/admin/resources')
def admin_resources():
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))
    
    try:
        resources = dbhelper.get_resources()
        return render_template('admin_resources.html', resources=resources)
    except Exception as e:
        print(f"Error: {e}")
        flash('Error loading resources', 'error')
        return redirect(url_for('admin_resources'))

@app.route('/upload_resource', methods=['POST'])
def upload_resource():
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))
    
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'})
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'resources', filename)
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            file.save(file_path)
            
            title = request.form.get('title', filename)
            description = request.form.get('description', '')
            file_type = get_file_type(filename)
            
            dbhelper.add_resource(title, description, filename, file_type)
            return jsonify({'success': True, 'message': 'Resource uploaded successfully'})
        else:
            return jsonify({'error': 'File type not allowed'})
    except Exception as e:
        print(f"Error: {e}")
        return jsonify({'error': 'Error uploading file'})

@app.route('/download_resource/<filename>')
def download_resource(filename):
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))   
    
    try:
        return send_from_directory(os.path.join(app.config['UPLOAD_FOLDER'], 'resources'), filename, as_attachment=True)
    except Exception as e:
        print(f"Error: {e}")
        flash('Error downloading file', 'error')
        return redirect(url_for('resources'))

def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'jpg', 'jpeg', 'png', 'gif', 'mp4', 'mov', 'avi'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_type(filename):
    extension = filename.rsplit('.', 1)[1].lower()
    if extension in ['pdf']:
        return 'pdf'
    elif extension in ['doc', 'docx', 'txt']:
        return 'document'
    elif extension in ['jpg', 'jpeg', 'png', 'gif']:
        return 'image'
    elif extension in ['mp4', 'mov', 'avi']:
        return 'video'
    return 'other'

@app.route('/delete_resource/<int:resource_id>', methods=['DELETE'])
def delete_resource(resource_id):
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({'error': 'Unauthorized access'}), 403

    try:
        resource = dbhelper.get_resource_by_id(resource_id)
        if not resource:
            return jsonify({'error': 'Resource not found'}), 404

        # Delete the file from the filesystem
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'resources', resource['filename'])
        if os.path.exists(file_path):
            os.remove(file_path)

        # Delete from the database
        dbhelper.delete_resource(resource_id)

        return jsonify({'success': True, 'message': 'Resource deleted successfully'})
    except Exception as e:
        print(f"Error deleting resource: {e}")
        return jsonify({'error': 'Failed to delete resource'}), 500

@app.route('/admin/computer_controller')
def admin_computer_controller():
    if 'username' not in session or session.get('role') != 'admin':
        flash('Please log in as admin first.', 'error')
        return redirect(url_for('login'))
    
    return render_template('admin_computer_controller.html')

@app.route('/admin/get_computers')
def get_computers():
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        laboratory = request.args.get('laboratory', '')
        status = request.args.get('status', '')
        search = request.args.get('search', '')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Base query
        query = """
            SELECT 
                c.id,
                c.computer_number,
                c.laboratory,
                c.status,
                u.firstname || ' ' || u.lastname as current_user,
                r.reservation_date || ' ' || r.start_time as reservation_time
            FROM computers c
            LEFT JOIN current_sitins cs ON c.id = cs.computer_id
            LEFT JOIN users u ON cs.user_id = u.id
            LEFT JOIN reservations r ON c.id = r.computer_id 
                AND r.status = 'approved' 
                AND r.reservation_date = date('now')
                AND r.start_time <= time('now')
                AND r.end_time > time('now')
            WHERE 1=1
        """
        params = []
        
        # Add filters
        if laboratory:
            query += " AND c.laboratory = ?"
            params.append(laboratory)
        
        if status:
            query += " AND c.status = ?"
            params.append(status)
        
        if search:
            query += " AND c.computer_number LIKE ?"
            params.append(f'%{search}%')
        
        query += " ORDER BY c.laboratory, c.computer_number"
        
        cursor.execute(query, params)
        computers = cursor.fetchall()
        
        # Format the results
        formatted_computers = []
        for computer in computers:
            formatted_computers.append({
                'id': computer[0],
                'computer_number': computer[1],
                'laboratory': computer[2],
                'status': computer[3],
                'current_user': computer[4],
                'reservation_time': computer[5]
            })
        
        conn.close()
        return jsonify({'computers': formatted_computers})
    except Exception as e:
        print(f"Error in get_computers: {e}")
        return jsonify({'error': 'Error getting computers'}), 500

@app.route('/admin/update_computer_status', methods=['POST'])
def update_computer_status():
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        data = request.get_json()
        computer_id = data.get('computer_id')
        status = data.get('status')
        
        if not computer_id or not status:
            return jsonify({'success': False, 'message': 'Missing required parameters'})
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Update computer status
        cursor.execute("""
            UPDATE computers
            SET status = ?
            WHERE id = ?
        """, (status, computer_id))
        
        # If setting to in_use, create a current_sitin record with admin as the user
        if status == 'in_use':
            # Get admin user ID from admins table
            cursor.execute("SELECT admin_id FROM admins WHERE username = ?", (session['username'],))
            admin_result = cursor.fetchone()
            
            if not admin_result:
                conn.close()
                return jsonify({'success': False, 'message': 'No admin user found in database'})
            
            admin_id = admin_result[0]
            
            cursor.execute("""
                INSERT INTO current_sitins (user_id, time_in, purpose, laboratory, computer_id)
                VALUES (?, datetime('now'), 'Walk-in (Admin)', (SELECT laboratory FROM computers WHERE id = ?), ?)
            """, (admin_id, computer_id, computer_id))
        
        conn.commit()
        conn.close()
        
        return jsonify({'success': True, 'message': 'Computer status updated successfully'})
    except Exception as e:
        print(f"Error in update_computer_status: {e}")
        return jsonify({'success': False, 'message': 'Error updating computer status'})

@app.route('/admin/make_computer_available', methods=['POST'])
def make_computer_available():
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        data = request.get_json()
        computer_id = data.get('computer_id')
        
        if not computer_id:
            return jsonify({'success': False, 'message': 'Missing computer ID'})
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Get the current sitin record
        cursor.execute("""
            SELECT id, user_id, time_in, laboratory
            FROM current_sitins
            WHERE computer_id = ?
        """, (computer_id,))
        
        sitin = cursor.fetchone()
        
        if sitin:
            # Add to sit_in_history
            cursor.execute("""
                INSERT INTO sit_in_history (user_id, time_in, time_out, purpose, laboratory, computer_id)
                VALUES (?, ?, datetime('now'), 'Walk-in', ?, ?)
            """, (sitin[1], sitin[2], sitin[3], computer_id))
            
            # Remove from current_sitins
            cursor.execute("DELETE FROM current_sitins WHERE id = ?", (sitin[0],))
        
        # Update computer status
        cursor.execute("""
            UPDATE computers
            SET status = 'available'
            WHERE id = ?
        """, (computer_id,))
        
        conn.commit()
        conn.close()
        
        return jsonify({'success': True, 'message': 'Computer made available successfully'})
    except Exception as e:
        print(f"Error in make_computer_available: {e}")
        return jsonify({'success': False, 'message': 'Error making computer available'})

@app.route('/get_notifications')
def get_notifications():
    if 'user_id' not in session:
        return jsonify({
            'success': False,
            'error': 'Unauthorized',
            'notifications': []
        }), 401
    
    try:
        notifications = dbhelper.get_user_notifications(session['user_id'])
        return jsonify({
            'success': True,
            'notifications': notifications
        })
    except Exception as e:
        print(f"Error getting notifications: {e}")
        return jsonify({
            'success': False,
            'error': str(e),
            'notifications': []
        }), 500

@app.route('/mark_notification_read/<int:notification_id>', methods=['POST'])
def mark_notification_read(notification_id):
    if 'user_id' not in session:
        return jsonify({
            'success': False,
            'error': 'Unauthorized'
        }), 401
    
    try:
        success = dbhelper.mark_notification_read(notification_id)
        if success:
            return jsonify({
                'success': True,
                'message': 'Notification marked as read'
            })
        else:
            return jsonify({
                'success': False,
                'error': 'Failed to mark notification as read'
            }), 400
    except Exception as e:
        print(f"Error marking notification as read: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/approve_reservation/<int:reservation_id>', methods=['POST'])
@login_required
def approve_reservation(reservation_id):
    try:
        if not current_user.is_admin:
            return jsonify({
                'success': False,
                'error': 'Unauthorized'
            }), 403
        
        print(f"Approving reservation {reservation_id}")  # Debug log
        reservation = Reservation.query.get_or_404(reservation_id)
        reservation.status = 'approved'
        
        # Create notification for the student
        notification = Notification(
            user_id=reservation.user_id,
            message=f"Your reservation request for Lab {reservation.laboratory} on PC{reservation.computer_id} has been approved"
        )
        print(f"Creating notification: {notification.message}")  # Debug log
        
        db.session.add(notification)
        db.session.commit()
        print("Successfully approved reservation and created notification")  # Debug log
        
        return jsonify({
            'success': True,
            'message': 'Reservation approved and notification sent'
        })
    except Exception as e:
        print(f"Error in approve_reservation: {str(e)}")  # Debug log
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/get_reservation_logs')
def get_reservation_logs():
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    
    try:
        filter_type = request.args.get('filter_type', 'all')
        filter_value = request.args.get('filter_value')
        
        logs = dbhelper.get_reservation_logs(filter_type, filter_value)
        return jsonify({'success': True, 'logs': logs})
    except Exception as e:
        print(f"Error getting reservation logs: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/get_student_points/<int:student_id>')
def get_student_points_route(student_id):
    if 'username' not in session or session.get('role') != 'admin':
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT points 
            FROM users 
            WHERE id = ?
        """, (student_id,))
        
        result = cursor.fetchone()
        points = result[0] if result and result[0] is not None else 0
        
        conn.close()
        return jsonify({'points': points})
    except Exception as e:
        print(f"Error getting student points: {e}")
        return jsonify({'error': 'Error fetching points'}), 500

@app.route('/view_schedule')
def view_schedule():
    if 'username' not in session:
        flash('Please log in first.', 'error')
        return redirect(url_for('login'))
    return render_template('view_schedule.html')

@app.route('/get_filtered_schedules', methods=['POST'])
def get_filtered_schedules():
    if 'username' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        data = request.get_json()
        laboratory = data.get('laboratory')
        day = data.get('day')
        time = data.get('time')
        
        schedules = dbhelper.get_filtered_schedules(laboratory, day, time)
        return jsonify({'schedules': schedules})
    except Exception as e:
        print(f"Error getting filtered schedules: {e}")
        return jsonify({'error': 'Error fetching schedules'}), 500

if __name__ == "__main__":
    # Initialize database tables
    try:
        dbhelper.create_tables()  # This will now only create tables if they don't exist
        print("Database tables initialized successfully")
    except Exception as e:
        print(f"Error initializing database: {e}")
    
    app.run(debug=True)
